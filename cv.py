from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import requests
from io import BytesIO
import json

def create_profile_document(data, exp):
    # Создание документа
    doc = Document()
    data = json.loads(data)
    exp = json.loads(exp)
    # Функция для вставки фото в ячейку таблицы
    def insert_image_in_cell(cell, image_url, width=1.5):
        response = requests.get(image_url)
        image_stream = BytesIO(response.content)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(width))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Функция для добавления текста в ячейку таблицы
    def insert_text_in_cell(cell, text, bold=False, font_size=12):
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Центрирование заголовка с именем
    name_paragraph = doc.add_paragraph(f"{data['firstName']} {data['lastName']}")
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.runs[0]
    name_run.font.size = Pt(20)

    # Добавление таблицы с одной строкой и двумя столбцами
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Добавление текстовых данных в левую ячейку (headline, industryName, locationName)
    text_cell = table.cell(0, 0)
    insert_text_in_cell(text_cell, data['experience'][0]['title'], bold=True, font_size=14)
    insert_text_in_cell(text_cell, data['industryName'], font_size=12)
    insert_text_in_cell(text_cell, f"{data['locationName']}, {data['geoCountryName']}", font_size=12)

    # Добавление изображения в правую ячейку
    image_cell = table.cell(0, 1)
    last_img_key = max((key for key in data if key.startswith('img_')), default=None)
    print
    try:
        insert_image_in_cell(image_cell, data['displayPictureUrl'] + data[last_img_key], width=1.5)
    except BaseException as e:
        print(e)

    # Секция "Summary"
    summary = data.get('summary', '')
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data['summary'])

    # Секция "Experience"
    doc.add_heading("Experience", level=1)
    for experience in exp:
        position = f"{experience['title']} at {experience['companyName']}"

        # Формируем период с учётом отсутствия месяца
        period = f"{experience['startDate']} - {experience['endDate']}"
        doc.add_paragraph(position, style='Heading 3')
        doc.add_paragraph(period)
        doc.add_paragraph(experience.get('description', ''))

    # Секция "Education"
    if len(data['education']) > 0:
        doc.add_heading("Education", level=1)
        for edu in data['education']:
            school = edu['schoolName']
            degree = edu.get('degreeName', '')
            field = edu.get('fieldOfStudy', '')
            try:
                start_year = edu['timePeriod']['startDate']['year']
            except KeyError:
                start_year = None
            try:
                end_year = edu['timePeriod']['endDate']['year']
            except KeyError:
                end_year = None
            # start_year = edu['timePeriod']['startDate']['year']
            # end_year = edu['timePeriod']['endDate']['year']
            if start_year and end_year:
                doc.add_paragraph(f"{school}, {degree} in {field} ({start_year} - {end_year})")
            else:
                doc.add_paragraph(f"{school}, {degree} in {field} ({start_year or end_year})")


    # Секция "Certifications"
    if len(data['certifications']) > 0:
        doc.add_heading("Certifications", level=1)
        for cert in data['certifications']:
            cert_name = cert['name']
            cert_authority = cert['authority']
            try:
                cert_date = f"{cert['timePeriod']['startDate']['month']}/{cert['timePeriod']['startDate']['year']}"
                doc.add_paragraph(f"{cert_name} - {cert_authority} ({cert_date})")
            except KeyError:
                doc.add_paragraph(f"{cert_name} - {cert_authority}")

    # Секция "Skills"
    doc.add_heading("Skills", level=1)
    skills_text = ', '.join(skill['name'] for skill in data['skills'])
    doc.add_paragraph(skills_text)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream, f"{data['firstName']}_{data['lastName']}.docx"
    # Сохранение файла

# Теперь можно вызывать функцию create_profile_document(data) из другого файла, передав туда нужные данные.
